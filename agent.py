# -*- coding: utf-8 -*-
"""
AI Overview Content Gap Agent
Justwords Screening Test - Stage 1
"""

import os, sys, json, time, glob, argparse, requests, re
from datetime import datetime
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except AttributeError:
        pass


# ──────────────────────────────────────────────
# SEO RANKINGS
# ──────────────────────────────────────────────

def fetch_seo_rankings(keyword: str, urls_to_check: list) -> dict:
    """
    Fetch organic ranking positions for the given keyword.
    Returns a dict with:
      - top10: list of {position, url, title, snippet}
      - client_rankings: {url -> position or None} for each url_to_check
    """
    serpapi_key = os.getenv("SERPAPI_KEY", "")
    serper_key  = os.getenv("SERPER_KEY", "")

    print("\n[SEO] Fetching organic ranking positions ...")

    try:
        if serpapi_key:
            results = _serpapi_organic(keyword, serpapi_key)
        elif serper_key:
            results = _serper_organic(keyword, serper_key)
        else:
            print("  [SEO] No SERP key found — skipping rankings.")
            return _empty_rankings(urls_to_check)
    except Exception as e:
        print(f"  [SEO] Could not fetch rankings: {e}")
        return _empty_rankings(urls_to_check)

    # Map each URL we care about to its rank
    client_rankings = {}
    for url in urls_to_check:
        if not url:
            continue
        domain = _extract_domain(url)
        match = next(
            (r for r in results if _extract_domain(r["url"]) == domain),
            None
        )
        if match:
            client_rankings[url] = match["position"]
            print(f"  [SEO] {url[:60]}... → rank #{match['position']}")
        else:
            client_rankings[url] = None
            print(f"  [SEO] {url[:60]}... → not in top 100")

    print(f"  [SEO] Top 10 organic results retrieved.")
    return {
        "top10": results[:10],
        "client_rankings": client_rankings,
    }


def _serpapi_organic(keyword, api_key):
    params = {
        "engine": "google", "q": keyword, "api_key": api_key,
        "gl": "in", "hl": "en", "num": 100,
    }
    resp = requests.get("https://serpapi.com/search", params=params, timeout=30)
    resp.raise_for_status()
    data = resp.json()
    return [
        {
            "position": r.get("position", i + 1),
            "url":      r.get("link", ""),
            "title":    r.get("title", ""),
            "snippet":  r.get("snippet", ""),
        }
        for i, r in enumerate(data.get("organic_results", []))
    ]


def _serper_organic(keyword, api_key):
    headers = {"X-API-KEY": api_key, "Content-Type": "application/json"}
    resp = requests.post(
        "https://google.serper.dev/search",
        headers=headers,
        json={"q": keyword, "gl": "in", "num": 100},
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    return [
        {
            "position": r.get("position", i + 1),
            "url":      r.get("link", ""),
            "title":    r.get("title", ""),
            "snippet":  r.get("snippet", ""),
        }
        for i, r in enumerate(data.get("organic", []))
    ]


def _extract_domain(url: str) -> str:
    try:
        from urllib.parse import urlparse
        return urlparse(url).netloc.lower().replace("www.", "")
    except Exception:
        return url


def _empty_rankings(urls_to_check):
    return {
        "top10": [],
        "client_rankings": {u: None for u in urls_to_check if u},
    }


# ──────────────────────────────────────────────
# SERP — AI Overview URLs
# ──────────────────────────────────────────────

def fetch_ai_overview_urls(keyword: str) -> list:
    serpapi_key = os.getenv("SERPAPI_KEY", "")
    serper_key  = os.getenv("SERPER_KEY", "")

    if serpapi_key:
        return _serpapi_fetch(keyword, serpapi_key)
    elif serper_key:
        return _serper_fetch(keyword, serper_key)
    else:
        raise EnvironmentError(
            "No search API key found. Set SERPAPI_KEY or SERPER_KEY in your .env file."
        )


def _serpapi_fetch(keyword, api_key):
    params = {
        "engine": "google", "q": keyword, "api_key": api_key,
        "gl": "in", "hl": "en", "num": 10,
    }
    resp = requests.get("https://serpapi.com/search", params=params, timeout=30)
    resp.raise_for_status()
    data = resp.json()
    urls = [
        s["link"]
        for s in data.get("ai_overview", {}).get("sources", [])
        if s.get("link")
    ]
    if not urls:
        urls = [r["link"] for r in data.get("organic_results", [])[:5] if r.get("link")]
    print(f"[SERP] Found {len(urls)} AI Overview URL(s) via SerpAPI.")
    return urls


def _serper_fetch(keyword, api_key):
    headers = {"X-API-KEY": api_key, "Content-Type": "application/json"}
    resp = requests.post(
        "https://google.serper.dev/search",
        headers=headers,
        json={"q": keyword, "gl": "in", "num": 10},
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    urls = [
        r["link"]
        for r in data.get("answerBox", {}).get("references", [])
        if r.get("link")
    ]
    if not urls:
        urls = [r["link"] for r in data.get("organic", [])[:5] if r.get("link")]
    print(f"[SERP] Found {len(urls)} AI Overview URL(s) via Serper.dev.")
    return urls


# ──────────────────────────────────────────────
# SCRAPER
# ──────────────────────────────────────────────

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0 Safari/537.36"
    )
}


def scrape_url(url: str, mock_html_path: str = None) -> dict:
    try:
        if mock_html_path:
            print(f"    [MOCK] Reading {os.path.basename(mock_html_path)} (no HTTP request)")
            with open(mock_html_path, "r", encoding="utf-8", errors="replace") as f:
                html = f.read()
        else:
            resp = requests.get(url, headers=HEADERS, timeout=20)
            resp.raise_for_status()
            html = resp.text

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = soup.get_text(separator=" ", strip=True)
        word_count = len(text.split())

        headings = {
            "h1": [h.get_text(strip=True) for h in soup.find_all("h1")],
            "h2": [h.get_text(strip=True) for h in soup.find_all("h2")],
            "h3": [h.get_text(strip=True) for h in soup.find_all("h3")],
        }

        has_faq = bool(
            soup.find(
                lambda t: t.name in ["h2", "h3", "div", "section"]
                and "faq" in (t.get_text() + " ".join(t.get("class", []))).lower()
            )
        )

        page_text_lower = text.lower()
        has_calculator = (
            any(w in page_text_lower for w in ["calculator", "calculate"])
            or bool(re.search(r'\bcompute\b', page_text_lower))
        )

        return {
            "url": url,
            "word_count": word_count,
            "headings": headings,
            "has_faq": has_faq,
            "has_table": bool(soup.find("table")),
            "has_numbered_list": bool(soup.find("ol")),
            "has_bullet_list": bool(soup.find("ul")),
            "has_calculator": has_calculator,
            "has_comparison": any(
                w in page_text_lower for w in ["vs", "versus", "compare", "comparison"]
            ),
            "full_text": text,
            "text_sample": text,
            "error": None,
        }
    except Exception as e:
        print(f"  [WARN] Could not scrape {url}: {e}")
        return {
            "url": url, "word_count": 0,
            "headings": {"h1": [], "h2": [], "h3": []},
            "has_faq": False, "has_table": False,
            "has_numbered_list": False, "has_bullet_list": False,
            "has_calculator": False, "has_comparison": False,
            "full_text": "", "text_sample": "", "error": str(e),
        }


# ──────────────────────────────────────────────
# LLM ANALYSIS
# ──────────────────────────────────────────────

def analyse_with_llm(keyword: str, client_data: dict, competitor_data: list) -> dict:
    groq_key      = os.getenv("GROQ_API_KEY", "")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")
    openai_key    = os.getenv("OPENAI_API_KEY", "")
    gemini_key    = os.getenv("GEMINI_API_KEY", "")
    ollama_model  = os.getenv("OLLAMA_MODEL", "")

    if groq_key:
        return _groq_analyse(keyword, client_data, competitor_data, groq_key)
    elif anthropic_key:
        return _claude_analyse(keyword, client_data, competitor_data, anthropic_key)
    elif openai_key:
        return _openai_analyse(keyword, client_data, competitor_data, openai_key)
    elif gemini_key:
        return _gemini_analyse(keyword, client_data, competitor_data, gemini_key)
    elif ollama_model:
        return _ollama_analyse(keyword, client_data, competitor_data, ollama_model)
    else:
        raise EnvironmentError(
            "No LLM API key found. Set one of the following in your .env file:\n"
            "  ANTHROPIC_API_KEY  - Claude (paid)\n"
            "  OPENAI_API_KEY     - GPT-4o-mini (paid)\n"
            "  GROQ_API_KEY       - Groq/Llama3 (free) -> https://console.groq.com\n"
            "  GEMINI_API_KEY     - Gemini (free) -> https://aistudio.google.com\n"
            "  OLLAMA_MODEL       - Local Ollama model, e.g. llama3 (free, offline)\n"
        )


def _build_prompt(keyword, client_data, competitor_data):
    SAMPLE_CHARS = 800

    comp_summaries = ""
    for i, c in enumerate(competitor_data, 1):
        h2s = "; ".join(c["headings"]["h2"][:8]) or "none"
        sample = c.get("full_text", c.get("text_sample", ""))[:SAMPLE_CHARS]
        comp_summaries += (
            f"\n--- Competitor {i}: {c['url']} ---\n"
            f"Word count: {c['word_count']}\n"
            f"H2 headings: {h2s}\n"
            f"FAQ: {c['has_faq']} | Tables: {c['has_table']} | "
            f"Numbered lists: {c['has_numbered_list']} | Bullets: {c['has_bullet_list']}\n"
            f"Calculator: {c['has_calculator']} | Comparison: {c['has_comparison']}\n"
            f"Content sample:\n{sample}\n"
        )

    client_h2s = "; ".join(client_data["headings"]["h2"][:8]) or "none"
    client_sample = client_data.get("full_text", client_data.get("text_sample", ""))[:SAMPLE_CHARS]
    client_summary = (
        f"URL: {client_data['url']}\n"
        f"Word count: {client_data['word_count']}\n"
        f"H2 headings: {client_h2s}\n"
        f"FAQ: {client_data['has_faq']} | Tables: {client_data['has_table']} | "
        f"Numbered lists: {client_data['has_numbered_list']} | Bullets: {client_data['has_bullet_list']}\n"
        f"Calculator: {client_data['has_calculator']} | Comparison: {client_data['has_comparison']}\n"
        f"Content sample:\n{client_sample}\n"
    )

    return f"""You are an expert SEO content strategist.

KEYWORD: "{keyword}"

CLIENT ARTICLE:
{client_summary}

AI OVERVIEW COMPETITOR ARTICLES:
{comp_summaries}

Your task:
1. Identify SPECIFIC content gaps: topics, sections, formats, or data points the client article is missing.
2. Provide 3-5 ACTIONABLE recommendations. Each must be concrete (e.g. "Add an FAQ answering X, Y, Z" not "Add an FAQ").
3. Give a brief content format comparison.

Respond ONLY in this JSON format (no markdown fences):
{{
  "format_comparison": {{
    "client_word_count": <int>,
    "avg_competitor_word_count": <int>,
    "client_has_faq": <bool>,
    "competitors_with_faq": <int>,
    "client_has_table": <bool>,
    "competitors_with_table": <int>,
    "client_has_numbered_list": <bool>,
    "client_has_bullet_list": <bool>
  }},
  "content_gaps": [
    {{"gap": "<short title>", "detail": "<specific explanation>"}}
  ],
  "recommendations": [
    {{"title": "<action title>", "detail": "<concrete steps>"}}
  ],
  "executive_summary": "<2-3 sentence plain-English summary>"
}}"""


def _build_prompt_short(keyword, client_data, competitor_data):
    comp_lines = ""
    for i, c in enumerate(competitor_data, 1):
        h2s = "; ".join(c["headings"]["h2"][:4]) or "none"
        comp_lines += (
            f"Competitor {i}: words={c['word_count']}, faq={c['has_faq']}, "
            f"table={c['has_table']}, h2s=[{h2s}]\n"
        )

    return f'''You are an SEO content strategist. Respond ONLY in valid JSON, no markdown.

KEYWORD: "{keyword}"
CLIENT: words={client_data["word_count"]}, faq={client_data["has_faq"]}, table={client_data["has_table"]}
COMPETITORS:
{comp_lines}

Return this exact JSON structure:
{{
  "format_comparison": {{
    "client_word_count": {client_data["word_count"]},
    "avg_competitor_word_count": 0,
    "client_has_faq": {str(client_data["has_faq"]).lower()},
    "competitors_with_faq": 0,
    "client_has_table": {str(client_data["has_table"]).lower()},
    "competitors_with_table": 0,
    "client_has_numbered_list": {str(client_data["has_numbered_list"]).lower()},
    "client_has_bullet_list": {str(client_data["has_bullet_list"]).lower()}
  }},
  "content_gaps": [
    {{"gap": "Gap 1 title", "detail": "Specific explanation of what is missing"}},
    {{"gap": "Gap 2 title", "detail": "Specific explanation of what is missing"}},
    {{"gap": "Gap 3 title", "detail": "Specific explanation of what is missing"}}
  ],
  "recommendations": [
    {{"title": "Action 1", "detail": "Concrete steps to take"}},
    {{"title": "Action 2", "detail": "Concrete steps to take"}},
    {{"title": "Action 3", "detail": "Concrete steps to take"}}
  ],
  "executive_summary": "2-3 sentence plain English summary of the key gaps and recommendations."
}}'''


def _parse_llm_json(raw: str, source: str) -> dict:
    cleaned = raw.strip().replace("```json", "").replace("```", "").strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        print(f"  [WARN] {source} returned malformed JSON: {e}")
        print(f"  [DEBUG] Raw output (first 500 chars):\n{cleaned[:500]}")
        raise RuntimeError(
            f"LLM ({source}) did not return valid JSON. "
            "Check the debug output above and retry."
        ) from e


def _claude_analyse(keyword, client_data, competitor_data, api_key):
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-sonnet-4-5", max_tokens=2000,
        messages=[{"role": "user", "content": _build_prompt(keyword, client_data, competitor_data)}],
    )
    return _parse_llm_json(msg.content[0].text, "Claude")


def _openai_analyse(keyword, client_data, competitor_data, api_key):
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": _build_prompt(keyword, client_data, competitor_data)}],
        max_tokens=2000,
        response_format={"type": "json_object"},
    )
    return _parse_llm_json(resp.choices[0].message.content, "OpenAI")


def _groq_analyse(keyword, client_data, competitor_data, api_key):
    from openai import OpenAI
    client = OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")

    for attempt, prompt in enumerate([
        _build_prompt(keyword, client_data, competitor_data),
        _build_prompt_short(keyword, client_data, competitor_data),
    ], 1):
        try:
            resp = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an SEO content strategist. "
                            "Always respond with valid JSON only. No markdown, no extra text."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                max_tokens=4000,
                temperature=0.1,
            )
            return _parse_llm_json(resp.choices[0].message.content, "Groq")
        except RuntimeError:
            if attempt == 2:
                raise
            print("  [INFO] Full prompt truncated — retrying with shorter prompt ...")


def _gemini_analyse(keyword, client_data, competitor_data, api_key):
    import urllib.request
    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-1.5-flash:generateContent?key={api_key}"
    )
    payload = json.dumps({
        "contents": [{"parts": [{"text": _build_prompt(keyword, client_data, competitor_data)}]}],
        "generationConfig": {"maxOutputTokens": 2000, "temperature": 0.2},
    }).encode()
    req = urllib.request.Request(url, data=payload,
                                 headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=60) as resp:
        data = json.loads(resp.read())
    raw = data["candidates"][0]["content"]["parts"][0]["text"]
    return _parse_llm_json(raw, "Gemini")


def _ollama_analyse(keyword, client_data, competitor_data, model):
    payload = json.dumps({
        "model": model,
        "prompt": _build_prompt(keyword, client_data, competitor_data),
        "stream": False,
    }).encode()
    req = requests.post(
        "http://localhost:11434/api/generate",
        data=payload,
        headers={"Content-Type": "application/json"},
        timeout=120,
    )
    req.raise_for_status()
    raw = req.json().get("response", "")
    return _parse_llm_json(raw, f"Ollama({model})")


# ──────────────────────────────────────────────
# REPORT GENERATION
# ──────────────────────────────────────────────

def generate_report(keyword, ai_overview_urls, client_data,
                    competitor_data, analysis, seo_rankings, output_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1.2)

    def add_heading(text, level=1, color=(31, 73, 125)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def shaded_row(table, cells, shade=False):
        row = table.add_row()
        for cell, val in zip(row.cells, cells):
            cell.text = str(val)
            if shade:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E2F3")
                tcPr.append(shd)
        return row

    def yn(v):
        return "Yes" if v else "No"

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AI Overview Content Gap Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 73, 125)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"Keyword: {keyword}  |  Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    doc.add_paragraph()

    add_heading("Executive Summary")
    doc.add_paragraph(analysis.get("executive_summary", "N/A"))

    # ── SEO Rankings section ──────────────────────────────────────────────
    add_heading("SEO Rankings")

    client_rankings = seo_rankings.get("client_rankings", {})
    if client_rankings:
        add_heading("Client Article Ranking", 2)
        for url, pos in client_rankings.items():
            rank_str = f"#{pos}" if pos else "Not in top 100"
            doc.add_paragraph(f"{url}\n  → Organic position: {rank_str}")

    top10 = seo_rankings.get("top10", [])
    if top10:
        add_heading("Top 10 Organic Results", 2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        shaded_row(tbl, ["Position", "URL", "Title"], shade=True)
        for r in top10:
            shaded_row(tbl, [
                f"#{r['position']}",
                r['url'][:60] + ("…" if len(r['url']) > 60 else ""),
                r['title'][:60] + ("…" if len(r['title']) > 60 else ""),
            ])
        doc.add_paragraph()

    # ── AI Overview URLs ──────────────────────────────────────────────────
    add_heading("AI Overview Source URLs")
    for i, url in enumerate(ai_overview_urls, 1):
        doc.add_paragraph(f"{i}. {url}", style="List Number")

    # ── Format analysis ───────────────────────────────────────────────────
    add_heading("Content Format Analysis")
    fc = analysis.get("format_comparison", {})
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    shaded_row(tbl, ["Dimension", "Client Article", "AI Overview Sources"], shade=True)
    for row_data in [
        ("Word Count",
         fc.get("client_word_count", client_data["word_count"]),
         fc.get("avg_competitor_word_count", "N/A")),
        ("FAQ Section",
         yn(fc.get("client_has_faq", client_data["has_faq"])),
         f"{fc.get('competitors_with_faq', '?')}/{len(competitor_data)} sources"),
        ("Tables",
         yn(fc.get("client_has_table", client_data["has_table"])),
         f"{fc.get('competitors_with_table', '?')}/{len(competitor_data)} sources"),
        ("Numbered Lists", yn(client_data["has_numbered_list"]), "See per-source below"),
        ("Bullet Lists",   yn(client_data["has_bullet_list"]),   "See per-source below"),
    ]:
        shaded_row(tbl, row_data)
    doc.add_paragraph()

    add_heading("Per-Source Format Detail", 2)
    for cd in competitor_data:
        if cd["error"]:
            doc.add_paragraph(f"- {cd['url']} -- could not scrape ({cd['error']})")
            continue
        h2s = ", ".join(cd["headings"]["h2"][:5]) or "none found"
        doc.add_paragraph(
            f"- {cd['url']}\n"
            f"  Words: {cd['word_count']} | FAQ: {yn(cd['has_faq'])} | "
            f"Table: {yn(cd['has_table'])} | Calculator: {yn(cd['has_calculator'])}\n"
            f"  Key H2s: {h2s}"
        )

    # ── Gaps & Recommendations ────────────────────────────────────────────
    add_heading("Identified Content Gaps")
    for gap in analysis.get("content_gaps", []):
        p = doc.add_paragraph(style="List Bullet")
        gap_run = p.add_run(gap.get("gap", "") + ": ")
        gap_run.bold = True
        p.add_run(gap.get("detail", ""))

    add_heading("Actionable Recommendations")
    for i, rec in enumerate(analysis.get("recommendations", []), 1):
        add_heading(f"{i}. {rec.get('title', '')}", 2, color=(31, 73, 125))
        doc.add_paragraph(rec.get("detail", ""))

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "Report generated by AI Overview Content Gap Agent."
    ).italic = True
    doc.save(output_path)
    print(f"[REPORT] Saved -> {output_path}")


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="AI Overview Content Gap Agent")
    parser.add_argument("--keyword",      required=True,
                        help='Search keyword, e.g. "best term insurance plan India"')
    parser.add_argument("--client-url",   required=False, default="",
                        help="Client article URL to analyse (optional)")
    parser.add_argument("--output",       default="gap_report.docx",
                        help="Output .docx path (default: gap_report.docx)")
    parser.add_argument("--mock-html-dir", default=None,
                        help=(
                            "Path to a folder of saved HTML files. Use instead of live "
                            "scraping when rate-limited, offline, or demoing. "
                            "Name files: competitor_1.html, competitor_2.html ... and client.html."
                        ))
    args = parser.parse_args()

    output_path = args.output
    if not output_path.lower().endswith(".docx"):
        output_path += ".docx"
        print(f"[INFO] Output path adjusted to: {output_path}")

    print(f"\n[Keyword] : {args.keyword}")
    print(f"[Client]  : {args.client_url if args.client_url else '(none)'}")
    if args.mock_html_dir:
        print(f"[Mock]    : reading HTML from '{args.mock_html_dir}' (no live scraping)")
    print()

    # ── Step 1: AI Overview URLs ──────────────────────────────────────────
    if args.mock_html_dir:
        print("Step 1/5 - Mock mode: skipping live SERP, building URLs from file names ...")
        comp_files = sorted(glob.glob(os.path.join(args.mock_html_dir, "competitor_*.html")))
        if not comp_files:
            print("  ERROR: No competitor_*.html files found in --mock-html-dir.")
            sys.exit(1)
        ai_urls = [
            "https://mock-source-" + os.path.splitext(os.path.basename(f))[0] + ".example.com"
            for f in comp_files
        ]
    else:
        print("Step 1/5 - Fetching AI Overview URLs ...")
        ai_urls = fetch_ai_overview_urls(args.keyword)
        if not ai_urls:
            print("  No AI Overview URLs found. Try a different keyword or use a VPN set to India.")
            sys.exit(0)
        comp_files = [None] * len(ai_urls)

    # ── Step 2: SEO Rankings ──────────────────────────────────────────────
    print("\nStep 2/5 - Fetching SEO rankings ...")
    urls_to_rank = [args.client_url] + ai_urls if args.client_url else ai_urls
    seo_rankings = fetch_seo_rankings(args.keyword, urls_to_rank)

    # ── Step 3: Scrape competitors ────────────────────────────────────────
    print(f"\nStep 3/5 - Processing {len(ai_urls)} competitor page(s) ...")
    competitor_data = []
    for url, fpath in zip(ai_urls, comp_files):
        print(f"  -> {url}")
        competitor_data.append(scrape_url(url, mock_html_path=fpath))
        if fpath is None:
            time.sleep(1)

    # ── Step 4: Scrape client ─────────────────────────────────────────────
    if args.client_url:
        print(f"\nStep 4/5 - Processing client URL ...")
        print(f"  -> {args.client_url}")
        if args.mock_html_dir:
            client_mock = os.path.join(args.mock_html_dir, "client.html")
            if not os.path.exists(client_mock):
                print(f"  ERROR: Expected client.html in {args.mock_html_dir} but not found.")
                sys.exit(1)
            client_data = scrape_url(args.client_url, mock_html_path=client_mock)
        else:
            time.sleep(1)
            client_data = scrape_url(args.client_url)
    else:
        print("\nStep 4/5 - No client URL provided, skipping ...")
        client_data = {
            "url": "", "word_count": 0,
            "headings": {"h1": [], "h2": [], "h3": []},
            "has_faq": False, "has_table": False,
            "has_numbered_list": False, "has_bullet_list": False,
            "has_calculator": False, "has_comparison": False,
            "full_text": "", "text_sample": "", "error": None,
        }

    # ── Step 5: LLM analysis ──────────────────────────────────────────────
    print("\nStep 5/5 - Running LLM gap analysis ...")
    analysis = analyse_with_llm(args.keyword, client_data, competitor_data)

    # ── Outputs ───────────────────────────────────────────────────────────
    json_path = os.path.splitext(output_path)[0] + ".json"

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({
            "keyword": args.keyword,
            "client_url": args.client_url,
            "ai_overview_urls": ai_urls,
            "seo_rankings": seo_rankings,
            "analysis": analysis,
            "generated_at": datetime.now().isoformat(),
        }, f, indent=2, ensure_ascii=False)
    print(f"[JSON]   Saved -> {json_path}")

    generate_report(
        keyword=args.keyword,
        ai_overview_urls=ai_urls,
        client_data=client_data,
        competitor_data=competitor_data,
        analysis=analysis,
        seo_rankings=seo_rankings,
        output_path=output_path,
    )

    print("\n[DONE]")
    print(f"   -> {output_path}")
    print(f"   -> {json_path}\n")


if __name__ == "__main__":
    main()
